"""Save generated documentation locally as Markdown, DOCX, and PDF (Task 24).

Writes under GeneratedDocs/<ProjectName>/, alongside the diagrams (as .mmd Mermaid source,
since neither python-docx nor reportlab can render Mermaid without a headless browser/Node
dependency this app doesn't otherwise need) and the resource inventory (as .csv, in addition
to being embedded as a table in every export format).
"""
import csv
import logging
import re
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document as DocxDocument
from docx.shared import Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from docgen.schemas import ProjectDocumentation

logger = logging.getLogger(__name__)

_GENERATED_DOCS_ROOT = Path("GeneratedDocs")

_SECTIONS: List[tuple] = [
    ("1. Executive Summary", "executive_summary"),
    ("2. Project Overview", "project_overview"),
    ("3. Architecture Overview", "architecture_overview"),
    ("4. Technology Stack", "technology_stack"),
    ("5. Repository Structure", "repository_structure"),
    ("6. Azure Infrastructure", "azure_infrastructure"),
    ("7. AKS Deployment", "aks_deployment"),
    ("8. App Service Deployment", "app_service_deployment"),
    ("9. CI/CD Pipeline", "cicd_pipeline"),
    ("10. Configuration", "configuration"),
    ("11. Networking", "networking"),
    ("12. Monitoring & Logging", "monitoring_logging"),
    ("13. Security Overview", "security_overview"),
    ("14. Resource Inventory", "resource_inventory_summary"),
    ("15. Deployment Flow", "deployment_flow"),
    ("16. Dependencies", "dependencies"),
    ("17. Troubleshooting Guide", "troubleshooting_guide"),
    ("19. Appendix", "appendix"),
]


def sanitize_project_folder_name(name: str) -> str:
    """A filesystem-safe folder name for GeneratedDocs/<ProjectName>/."""
    safe = re.sub(r"[^A-Za-z0-9_\-]+", "_", (name or "project").strip())
    return safe.strip("_") or "project"


def project_output_dir(project_name: str) -> Path:
    out_dir = _GENERATED_DOCS_ROOT / sanitize_project_folder_name(project_name)
    out_dir.mkdir(parents=True, exist_ok=True)
    return out_dir


def save_markdown(markdown_text: str, out_dir: Path) -> Path:
    path = out_dir / "documentation.md"
    path.write_text(markdown_text, encoding="utf-8")
    return path


def save_diagrams(dependency_diagram: str, deployment_flow_diagram: str, out_dir: Path) -> Dict[str, Path]:
    dep_path = out_dir / "architecture_diagram.mmd"
    flow_path = out_dir / "deployment_diagram.mmd"
    dep_path.write_text(dependency_diagram, encoding="utf-8")
    flow_path.write_text(deployment_flow_diagram, encoding="utf-8")
    return {"architecture_diagram": dep_path, "deployment_diagram": flow_path}


def save_resource_inventory_csv(rows: List[List[str]], out_dir: Path) -> Path:
    path = out_dir / "resource_inventory.csv"
    with path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.writer(handle)
        writer.writerows(rows)
    return path


def save_docx(
    document: ProjectDocumentation,
    inventory_rows: List[List[str]],
    metadata: Dict[str, Any],
    out_dir: Path,
) -> Path:
    """Render the documentation as a .docx file. Mermaid diagrams are embedded as their raw
    source (labelled) rather than as images - see module docstring."""
    doc = DocxDocument()
    doc.add_heading(f"{document.project_name} - Project Documentation", level=0)
    doc.add_paragraph(f"Generated {metadata.get('generated_at', 'Unknown')} from live Azure/GitLab/AKS data.")

    for title, field_name in _SECTIONS:
        doc.add_heading(title, level=1)
        doc.add_paragraph(getattr(document, field_name, "") or "Not available.")
        if field_name == "architecture_overview":
            doc.add_heading("Infrastructure Dependency Diagram (Mermaid source)", level=2)
            _add_monospace_paragraph(doc, metadata.get("dependency_diagram", ""))
        if field_name == "cicd_pipeline":
            doc.add_heading("Deployment Flow Diagram (Mermaid source)", level=2)
            _add_monospace_paragraph(doc, metadata.get("deployment_flow_diagram", ""))
        if field_name == "resource_inventory_summary" and len(inventory_rows) > 1:
            _add_docx_table(doc, inventory_rows)

    doc.add_heading("18. AI Recommendations", level=1)
    if document.ai_recommendations:
        for item in document.ai_recommendations:
            doc.add_paragraph(item, style="List Bullet")
    else:
        doc.add_paragraph("No recommendations were generated.")

    doc.add_heading("19. Appendix", level=1)
    doc.add_paragraph(document.appendix or "Not available.")

    if document.data_completeness_notes:
        doc.add_heading("Data Completeness Notes", level=1)
        for note in document.data_completeness_notes:
            doc.add_paragraph(note, style="List Bullet")

    path = out_dir / "documentation.docx"
    doc.save(str(path))
    return path


def _add_monospace_paragraph(doc: DocxDocument, text: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text or "")
    run.font.name = "Consolas"
    run.font.size = Pt(8)


def _add_docx_table(doc: DocxDocument, rows: List[List[str]]) -> None:
    header, data_rows = rows[0], rows[1:]
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Light Grid Accent 1"
    for cell, value in zip(table.rows[0].cells, header):
        cell.text = str(value)
    for row in data_rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            cell.text = str(value)


def save_pdf(
    document: ProjectDocumentation,
    inventory_rows: List[List[str]],
    metadata: Dict[str, Any],
    out_dir: Path,
) -> Path:
    """Render the documentation as a .pdf file via reportlab (pure Python, no system deps)."""
    path = out_dir / "documentation.pdf"
    styles = getSampleStyleSheet()
    mono_style = ParagraphStyle("Mono", parent=styles["Code"], fontSize=6, leading=8)

    story = [
        Paragraph(f"{_escape(document.project_name)} - Project Documentation", styles["Title"]),
        Paragraph(f"Generated {_escape(metadata.get('generated_at', 'Unknown'))} from live Azure/GitLab/AKS data.", styles["Normal"]),
        Spacer(1, 12),
    ]

    for title, field_name in _SECTIONS:
        story.append(Paragraph(_escape(title), styles["Heading1"]))
        story.append(Paragraph(_escape(getattr(document, field_name, "") or "Not available."), styles["Normal"]))
        story.append(Spacer(1, 6))
        if field_name == "architecture_overview":
            story.append(Paragraph("Infrastructure Dependency Diagram (Mermaid source):", styles["Heading3"]))
            story.append(Paragraph(_escape(metadata.get("dependency_diagram", "")).replace("\n", "<br/>"), mono_style))
            story.append(Spacer(1, 6))
        if field_name == "cicd_pipeline":
            story.append(Paragraph("Deployment Flow Diagram (Mermaid source):", styles["Heading3"]))
            story.append(Paragraph(_escape(metadata.get("deployment_flow_diagram", "")).replace("\n", "<br/>"), mono_style))
            story.append(Spacer(1, 6))
        if field_name == "resource_inventory_summary" and len(inventory_rows) > 1:
            story.append(_build_pdf_table(inventory_rows))
            story.append(Spacer(1, 6))

    story.append(Paragraph("18. AI Recommendations", styles["Heading1"]))
    if document.ai_recommendations:
        for item in document.ai_recommendations:
            story.append(Paragraph(f"- {_escape(item)}", styles["Normal"]))
    else:
        story.append(Paragraph("No recommendations were generated.", styles["Normal"]))

    story.append(Paragraph("19. Appendix", styles["Heading1"]))
    story.append(Paragraph(_escape(document.appendix or "Not available."), styles["Normal"]))

    if document.data_completeness_notes:
        story.append(Paragraph("Data Completeness Notes", styles["Heading1"]))
        for note in document.data_completeness_notes:
            story.append(Paragraph(f"- {_escape(note)}", styles["Normal"]))

    SimpleDocTemplate(str(path), pagesize=LETTER).build(story)
    return path


def _escape(text: str) -> str:
    return (text or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _build_pdf_table(rows: List[List[str]]) -> Table:
    escaped = [[_escape(str(cell)) for cell in row] for row in rows]
    table = Table(escaped, repeatRows=1)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2E86AB")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTSIZE", (0, 0), (-1, -1), 6),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
    ]))
    return table


def export_all(
    document: ProjectDocumentation,
    markdown_text: str,
    dependency_diagram: str,
    deployment_flow_diagram: str,
    inventory_rows: List[List[str]],
    metadata: Dict[str, Any],
) -> Dict[str, str]:
    """Save Markdown, DOCX, PDF, diagrams, and the CSV inventory under
    GeneratedDocs/<ProjectName>/, returning {label: path} for every file written."""
    out_dir = project_output_dir(document.project_name)
    metadata = {**metadata, "dependency_diagram": dependency_diagram, "deployment_flow_diagram": deployment_flow_diagram}

    paths: Dict[str, Path] = {"markdown": save_markdown(markdown_text, out_dir)}
    try:
        paths["docx"] = save_docx(document, inventory_rows, metadata, out_dir)
    except Exception as exc:
        logger.error("DOCX export failed for %s: %s", document.project_name, exc)
    try:
        paths["pdf"] = save_pdf(document, inventory_rows, metadata, out_dir)
    except Exception as exc:
        logger.error("PDF export failed for %s: %s", document.project_name, exc)

    paths.update(save_diagrams(dependency_diagram, deployment_flow_diagram, out_dir))
    if len(inventory_rows) > 1:
        paths["resource_inventory_csv"] = save_resource_inventory_csv(inventory_rows, out_dir)

    return {label: str(path) for label, path in paths.items()}
